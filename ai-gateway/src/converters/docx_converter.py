"""
DOCX Document Converter
Converts Microsoft Word documents to structured text using python-docx and Polars
"""

import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches
import polars as pl

logger = logging.getLogger(__name__)

class DOCXConverter:
    """Converts DOCX documents to structured text with advanced analysis"""
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file format is supported"""
        return file_path.suffix.lower() in self.supported_formats
    
    def extract_text(self, file_path: Path, include_metadata: bool = True) -> Dict[str, Any]:
        """
        Extract text from DOCX file using python-docx with Polars analysis
        
        Args:
            file_path: Path to DOCX file
            include_metadata: Whether to include document metadata
            
        Returns:
            Dictionary containing extracted text, metadata, and structured data
        """
        try:
            doc = Document(file_path)
            
            # Extract metadata if requested
            metadata = {}
            if include_metadata:
                metadata = self._extract_metadata(doc, file_path)
            
            # Extract text from paragraphs
            paragraphs_data = []
            tables_data = []
            total_chars = 0
            
            # Process paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                char_count = len(text)
                total_chars += char_count
                
                para_data = {
                    'element_type': 'paragraph',
                    'element_number': i + 1,
                    'text': text,
                    'char_count': char_count,
                    'has_text': char_count > 0,
                    'style': paragraph.style.name if paragraph.style else 'Normal'
                }
                
                # Check for special formatting
                if paragraph.runs:
                    formatting = []
                    for run in paragraph.runs:
                        if run.bold:
                            formatting.append('bold')
                        if run.italic:
                            formatting.append('italic')
                        if run.underline:
                            formatting.append('underline')
                    para_data['formatting'] = formatting
                
                paragraphs_data.append(para_data)
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                table_data = {
                    'element_type': 'table',
                    'element_number': i + 1,
                    'text': table_text,
                    'char_count': len(table_text),
                    'has_text': len(table_text) > 0,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0
                }
                tables_data.append(table_data)
                total_chars += len(table_text)
            
            # Combine all elements
            all_elements = paragraphs_data + tables_data
            
            # Create structured output using Polars
            try:
                elements_df = pl.DataFrame(all_elements)
            except Exception as e:
                logger.warning(f"Failed to create Polars DataFrame: {e}")
                elements_df = None
            
            # Calculate statistics
            statistics = {
                'total_paragraphs': len(paragraphs_data),
                'total_tables': len(tables_data),
                'total_elements': len(all_elements),
                'total_characters': total_chars,
                'non_empty_elements': len([e for e in all_elements if e['has_text']]),
                'empty_elements': len([e for e in all_elements if not e['has_text']]),
                'extraction_method': 'python_docx_with_polars'
            }
            
            return {
                'metadata': metadata,
                'paragraphs': paragraphs_data,
                'tables': tables_data,
                'all_elements': all_elements,
                'elements_dataframe': elements_df,
                'statistics': statistics,
                'full_text': self._combine_element_text(all_elements),
                'status': 'success'
            }
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            return {
                'metadata': {'filename': file_path.name},
                'paragraphs': [],
                'tables': [],
                'all_elements': [],
                'elements_dataframe': pl.DataFrame(),
                'statistics': {'total_elements': 0, 'total_characters': 0, 'extraction_method': 'failed'},
                'full_text': '',
                'error': str(e),
                'status': 'error'
            }
    
    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata"""
        metadata = {
            'filename': file_path.name,
            'file_size': file_path.stat().st_size,
        }
        
        # Try to extract core properties
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
            })
        except Exception as e:
            logger.warning(f"Failed to extract core properties: {e}")
        
        return metadata
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a table"""
        table_texts = []
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_texts.append(cell_text)
            if any(row_texts):  # Only add non-empty rows
                table_texts.append(' | '.join(row_texts))
        return '\n'.join(table_texts)
    
    def _combine_element_text(self, elements_data: List[Dict]) -> str:
        """Combine text from all elements into a single string"""
        texts = []
        for element in elements_data:
            if element.get('text') and element['has_text']:
                element_type = element.get('element_type', 'element')
                element_num = element.get('element_number', '')
                texts.append(f"--- {element_type.title()} {element_num} ---\n{element['text']}\n")
        return '\n'.join(texts)
    
    def analyze_structure(self, file_path: Path) -> Dict[str, Any]:
        """
        Analyze document structure and return insights using Polars
        """
        result = self.extract_text(file_path)
        if result['status'] != 'success' or result['elements_dataframe'] is None:
            return result
        
        df = result['elements_dataframe']
        
        # Analyze document structure
        analysis = {
            'content_distribution': {
                'paragraphs': len(result['paragraphs']),
                'tables': len(result['tables']),
                'elements_with_content': df.filter(pl.col('has_text')).height,
                'elements_without_content': df.filter(~pl.col('has_text')).height,
                'content_percentage': (df.filter(pl.col('has_text')).height / df.height * 100) if df.height > 0 else 0
            },
            'text_statistics': {
                'total_characters': df.select(pl.col('char_count').sum()).item(),
                'average_chars_per_element': df.select(pl.col('char_count').mean()).item(),
                'max_chars_per_element': df.select(pl.col('char_count').max()).item(),
                'min_chars_per_element': df.select(pl.col('char_count').min()).item(),
            }
        }
        
        # Analyze paragraph styles if available
        if 'style' in df.columns:
            try:
                style_counts = df.filter(pl.col('element_type') == 'paragraph').select('style').value_counts()
                analysis['style_distribution'] = style_counts.to_dict()
            except Exception as e:
                logger.warning(f"Failed to analyze styles: {e}")
        
        # Detect potential issues
        issues = []
        if analysis['content_distribution']['content_percentage'] < 80:
            issues.append("More than 20% of elements appear to be empty")
        
        if analysis['text_statistics']['average_chars_per_element'] < 50:
            issues.append("Very short average element length - document may be poorly structured")
        
        return {
            **result,
            'structure_analysis': analysis,
            'potential_issues': issues
        }